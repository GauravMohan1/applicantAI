import os

import openai
from flask import Flask, redirect, render_template, request, url_for, jsonify
import PyPDF2
from dotenv import load_dotenv, find_dotenv
from pathlib import Path
import mimetypes
from io import BytesIO
from docx import Document



app = Flask(__name__)
load_dotenv(Path('/etc/secrets/.env'))
openai.api_key = os.getenv("OPENAI_API_KEY")

@app.route("/")
def index():
    return render_template("index.html")
    
@app.route("/app", methods=["POST"])
def get_letter():
    error_message = "Please upload a PDF or DOCX version of your resume"
    company = request.form["company"]
    industry = request.form["industry"]
    request_file = request.files['resume']
    file_name = request_file.filename
    file_type = mimetypes.guess_type(file_name)[0]
    if file_type == 'application/pdf':
        resume = extract_resume_pdf(request_file)
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        resume = extract_resume_docx(request_file)
    else:
        return render_template('index.html', result=error_message)
    description = request.form['job']
    job_desc = extract_job(description)
    
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=generate_prompt(resume, company, job_desc, industry),
        temperature=0.75,
        max_tokens = 2300,
        n=2
    )

    cover_letter = response.choices[0].text
    return render_template('index.html', result=cover_letter)
   
def generate_prompt(resume,company, job_desc, industry):
    prompt = {'Finance': "The first paragraph should highlight why the company is interesting to the applicant based off what the company does, company culture, and the company's mission. The second paragraph should align the role responsibilities with the applicant's accomplishments. The third paragraph should mention the leadership abilities of the applicant and explain how that aligns with the company's mission and values.",
              'Business': "The first paragraph should highlight why the company is interesting to the applicant based off what the company does, company culture, and the company's mission. The second paragraph should highlight what the applicant has accomplished and how the applicant's internship and professional experience aligns with the role qualifications. The third paragraph should talk about the applicant's communication and leadership skills, including what the applicant accomplished in that leadership position, and explain how that fits into to the role and company.",
              'Product': "The first paragraph should highlight why the company is interesting to the applicant based off what the company does, company culture, and the company's mission. The second paragraph should highlight the various products the applicant has worked on and his or her impact on improving the product and align these experiences with the company's product. The third paragraph should talk about the applicant's soft skills, such as communication, time management, and creativity and why that will help him or her succeed on the job.",
              'HR': "The first paragraph should highlight why the company is interesting to the applicant based off what the company does, company culture, and the company's mission. The second paragraph should highlight the applicant's initiative to help others through great communication and why that leadership is important in this role. The third paragraph should highlight the applicant's approach to resolving conflict and how that aligns with the role responsibilities and company values.",
              'Engineering': "The first paragraph should highlight why the company is interesting to the applicant based off what the company does, company culture, and the company's mission. The second paragraph should highlight the applicant's impact created from their work experience and projects he or she developed. The third paragraph should mention how the role responsibilities aligns with the applicant's skills and experience why this makes them a perfect fit for the company.",
              'Manager': "The first paragraph should highlight why the company is interesting to the applicant based off what the company does, the company culture, and the company's mission. The second paragraph should highlight the types of teams the applicant has managed and their leadership style. The third paragraph should mention the accomplishements of the teams' the applicant has managed and how that fits into the role responsibilities and qualifications."}
    string = "Write a three paragraph cover letter for a job application using the information below."
    string += "\n"
    string += 'Company: ' + company.capitalize() + "\n"
    string += 'Resume: ' + resume + "\n"
    string += 'Job Description: ' + job_desc + "\n"
    string +=  prompt[industry]
    return string


def extract_resume_pdf(pdf_file):
    # Get the binary data of the PDF file from the request
    # Create a PDF object
    pdf = PyPDF2.PdfReader(pdf_file)

    # Iterate over every page in the PDF
    text = ""
    for page in range(len(pdf.pages)):
        # Extract the text from the current page
        text += pdf.pages[page].extract_text()
    # Return the extracted text
# Apply text on GPT-3 Summarization
    response = openai.Completion.create(
                model="text-davinci-003",
                prompt="""Summarize the text below into a JSON with exactly the following structure {basic_info: {first_name, last_name, full_name, email, phone_number, location, portfolio_website_url, linkedin_url, github_main_page_url, university, education_level (BS, MS, or PhD), graduation_year, graduation_month, majors, GPA}, work_experience: [{job_title, company, location, duration, job_summary}], leadership_experience:[{role, description}], project_experience:[{project_name, project_discription}]}
""" + '\n' + text,
                temperature = 0.0,
                max_tokens = 2000
            )

    response_text = response['choices'][0]['text'].strip()

    return response_text

def extract_resume_docx(docx_file):
    file_bytes = BytesIO()
    file_bytes.write(docx_file.read())
    # Seek to the beginning of the BytesIO object
    file_bytes.seek(0)
    # Read the file using the Document object from python-docx
    document = Document(file_bytes)
    # Extract the contents of the document
    text = ''
    for paragraph in document.paragraphs:
        text += paragraph.text
    
    response = openai.Completion.create(
                model="text-davinci-003",
                prompt="""Summarize the text below into a JSON with exactly the following structure {basic_info: {first_name, last_name, full_name, email, phone_number, location, portfolio_website_url, linkedin_url, github_main_page_url, university, education_level (BS, MS, or PhD), graduation_year, graduation_month, majors, GPA}, work_experience: [{job_title, company, location, duration, job_summary}], leadership_experience:[{role, description}], project_experience:[{project_name, project_discription}]}
""" + '\n' + text,
                temperature = 0.0,
                max_tokens = 2000
            )

    response_text = response['choices'][0]['text'].strip()

    return response_text


def extract_job(description):
    response_text = ''
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt="""Summarize the text below into a JSON with exactly the following structure {basic_info: {company description}, role_description: [{job_title, responsibilities}], role_qualifications:[{qualifications}]}
""" + '\n' + description,
        temperature = 0.0,
        max_tokens = 1000
    )
    response_text += response['choices'][0]['text'].strip()

    return response_text



@app.route("/resume", methods=["GET"])
def resume():
    return render_template("form.html")

@app.route("/resume", methods=["POST"])
def get_recommendations():
    description = request.form['job']   
    job_desc = extract_job(description)
   
    error_message = "Please upload a PDF or DOCX version of your resume"
    request_file = request.files["resume"]
    file_name = request_file.filename
    file_type = mimetypes.guess_type(file_name)[0]
    if file_type == 'application/pdf':
        resume = extract_resume_pdf(request_file)
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        resume = extract_resume_docx(request_file)
    else:
        return render_template('form.html', result=error_message)

    prompt = "Provide recommendations on improving the following resume given the job description and seperate the recommendations into work experience, project experience, leadership experience. Recommendations should quote from the applicant's resume when explaining how to improve the match rate with the job description, improve word usage, and sentence structure."
    prompt += 'Resume: ' + resume + "\n"
    prompt += 'Job Description' + job_desc
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=prompt,
        temperature=0.8,
        max_tokens = 2485
    )
    response_text = response['choices'][0]['text'].strip()
    return render_template('form.html', result=response_text)


def extract_resume_pdf(pdf_file):
    # Get the binary data of the PDF file from the request
    # Create a PDF object
    pdf = PyPDF2.PdfReader(pdf_file)

    # Iterate over every page in the PDF
    text = ""
    for page in range(len(pdf.pages)):
        # Extract the text from the current page
        text += pdf.pages[page].extract_text()
    # Return the extracted text
# Apply text on GPT-3 Summarization
    response = openai.Completion.create(
                model="text-davinci-003",
                prompt="""Summarize the text below into a JSON with exactly the following structure {basic_info: {first_name, last_name, full_name, email, phone_number, location, portfolio_website_url, linkedin_url, github_main_page_url, university, education_level (BS, MS, or PhD), graduation_year, graduation_month, majors, GPA}, work_experience: [{job_title, company, location, duration, job_summary}], leadership_experience:[{role, description}], project_experience:[{project_name, project_discription}]}
""" + '\n' + text,
                temperature = 0.0,
                max_tokens = 2000
            )

    response_text = response['choices'][0]['text'].strip()

    return response_text

def extract_resume_docx(docx_file):
    file_bytes = BytesIO()
    file_bytes.write(docx_file.read())
    # Seek to the beginning of the BytesIO object
    file_bytes.seek(0)
    # Read the file using the Document object from python-docx
    document = Document(file_bytes)
    # Extract the contents of the document
    text = ''
    for paragraph in document.paragraphs:
        text += paragraph.text
    
    response = openai.Completion.create(
                model="text-davinci-003",
                prompt="""Summarize the text below into a JSON with exactly the following structure {basic_info: {first_name, last_name, full_name, email, phone_number, location, portfolio_website_url, linkedin_url, github_main_page_url, university, education_level (BS, MS, or PhD), graduation_year, graduation_month, majors, GPA}, work_experience: [{job_title, company, location, duration, job_summary}], leadership_experience:[{role, description}], project_experience:[{project_name, project_discription}]}
""" + '\n' + text,
                temperature = 0.0,
                max_tokens = 2000
            )

    response_text = response['choices'][0]['text'].strip()

    return response_text

def extract_job(description):
    response_text = ''
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt="""Summarize the text below into a JSON with exactly the following structure {basic_info: {company description}, role_description: [{job_title, responsibilities}], role_qualifications:[{qualifications, skills}]}
""" + '\n' + description,
        temperature = 0.0,
        max_tokens = 1000
    )
    response_text += response['choices'][0]['text'].strip()

    return response_text



@app.route("/editor", methods=["GET"])
def resume():
    return render_template("editor.html")

@app.route("/editor", methods=["POST"])
def get_section():
    description = request.form['job']   
    job_desc = extract_job(description)

    section = request.form['section']
    new_section = rewrite(section,job_desc)
    return render_template('form.html', result=new_section)

def extract_job(description):
    response_text = ''
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt="""Summarize the text below into a JSON with exactly the following structure {basic_info: {company description}, role_description: [{job_title, responsibilities}], role_qualifications:[{qualifications, skills}]}
""" + '\n' + description,
        temperature = 0.0,
        max_tokens = 1000
    )
    response_text += response['choices'][0]['text'].strip()

    return response_text

    
def rewrite(section,job_desc):
    string = "Rewrite the following experience section to more closely align with the job qualifications and responsibilities. These changes may include incorporating relevant skills, improving sentence structure, change wording and highlighting accomplishments. Maintain the same format\
        of the original section."
    string += "\n"
    string += 'Experience Section: ' + section + "\n"
    string += 'Job Description: ' + job_desc
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=string,
        temperature=0.9,
        max_tokens = 1000,
        n=2
    )
    return response['choices'][0]['text'].strip()